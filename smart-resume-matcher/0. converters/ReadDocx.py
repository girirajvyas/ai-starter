import docx
def getTextViaDocx_v1(filePath):
    textFile = []
    print('processing ' + filePath)
    document = docx.Document(filePath)
    for paragraph in document.paragraphs:
        textFile.append(paragraph.text)
    
    return '\n'.join(textFile)

from docx import Document
def getTextViaDocx_v2(path):
	document = Document(path)
	return '\n'.join([para.text for para in document.paragraphs])